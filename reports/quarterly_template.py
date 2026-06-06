"""
Quarterly competitive intelligence report generator.

Entry point
-----------
    from reports.quarterly_template import generate_quarterly_report
    docx_bytes = await generate_quarterly_report("tenant_001", "2024Q4")

Sections produced
-----------------
    Cover          — institution name, period, confidential notice
    1 Executive Summary        Claude, ~200 words
    2 Market Share Dashboard   colour-coded table (measured/modeled/estimated)
    3 Competitive Movements    Claude per geography, parallel calls
    4 Peer Benchmarking        table + Claude narrative
    5 Market Opportunities     Claude top-3 geographies
    6 Data Notes               sources, confidence counts, next release dates

Period formats accepted:  "2024Q4"  |  "2024"  (treated as Q4)

DB tables read (all read-only):
    monitored_geographies   — tenant's watched markets
    alerts                  — recent significant-change notifications
    institutions_quarterly  — NCUA institution-level data for peer selection
"""
from __future__ import annotations

import asyncio
import io
import os
from datetime import datetime
from typing import Optional

import anthropic
import pandas as pd
import sqlalchemy as sa
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from ..database import get_engine
from ..processing.market_share_engine import calculate_market_share

# ── Claude client ──────────────────────────────────────────────────────────────

_client = anthropic.AsyncAnthropic()
_MODEL  = "claude-opus-4-8"

# ── Metric config ──────────────────────────────────────────────────────────────

# (engine_metric, table_key, display_label, unit)
_SHARE_METRICS = [
    ("deposits",              "deposit_share",   "Deposit Market Share",   "pct"),
    ("loans",                 "loan_share",      "Loan Market Share",      "pct"),
    ("mortgage_originations", "mortgage_share",  "Mortgage Market Share",  "pct"),
    ("members",               "member_share",    "Member Market Share",    "pct"),
]

# Peer benchmarking table — 6 rows
_PEER_METRIC_ROWS = [
    ("deposit_share",   "Deposit Market Share",    "pct"),
    ("loan_share",      "Loan Market Share",       "pct"),
    ("mortgage_share",  "Mortgage Market Share",   "pct"),
    ("member_share",    "Member Market Share",     "pct"),
    ("deposit_growth",  "Deposit Share Chg (QoQ)", "pp"),
    ("total_branches",  "Total Branches",          "int"),
]

# ── Document palette ───────────────────────────────────────────────────────────

_COL_HEADER = "1E3A8A"   # deep-blue table/section header
_COL_GREEN  = "C6EFCE"   # above peer median
_COL_AMBER  = "FFEB9C"   # below peer median
_COL_ALT    = "EFF6FF"   # alternating row tint
_WHITE      = "FFFFFF"

_RGB_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
_RGB_BRAND  = RGBColor(0x21, 0x63, 0xEB)
_RGB_DARK   = RGBColor(0x0F, 0x17, 0x2A)
_RGB_MUTED  = RGBColor(0x94, 0xA3, 0xB8)
_RGB_RED    = RGBColor(0xDC, 0x26, 0x26)

_MAX_PEERS   = 10
_MAX_GEO     = 12    # cap renders in the document

# ── Period helpers ─────────────────────────────────────────────────────────────

def _parse_period(period: str) -> tuple[int, int]:
    p = period.upper()
    if "Q" in p:
        idx = p.index("Q")
        return int(p[:idx]), int(p[idx + 1])
    return int(p), 4

def _fmt_period(year: int, q: int) -> str:
    return f"{year}Q{q}"

def _prior_quarter(year: int, q: int) -> str:
    return _fmt_period(year - 1, 4) if q == 1 else _fmt_period(year, q - 1)

def _yoy_period(year: int, q: int) -> str:
    return _fmt_period(year - 1, q)

def _human_period(period: str) -> str:
    p = period.upper()
    if "Q" in p:
        idx = p.index("Q")
        return f"Q{p[idx+1]} {p[:idx]}"
    return period

# ── Database helpers ───────────────────────────────────────────────────────────

async def _fetch_monitored_geos(tenant_id: str, engine: sa.engine.Engine) -> list[dict]:
    def _q(eng):
        with eng.connect() as conn:
            rows = conn.execute(
                sa.text("""
                    SELECT geography_type, geography_id, geography_label,
                           metric, own_institution_id, direct_competitor_ids
                    FROM   monitored_geographies
                    WHERE  tenant_id = :tid
                    ORDER  BY geography_label
                """),
                {"tid": tenant_id},
            ).fetchall()
        return [dict(r._mapping) for r in rows]
    return await asyncio.to_thread(_q, engine)


async def _fetch_own_institution(own_id: str, period: str, engine: sa.engine.Engine) -> dict:
    """Look up name, state, assets from institutions_quarterly for the own CU."""
    year, q = _parse_period(period)
    # Try exact period first, fall back to latest available
    def _q(eng):
        with eng.connect() as conn:
            row = conn.execute(
                sa.text("""
                    SELECT institution_name, state, total_assets, number_of_branches
                    FROM   institutions_quarterly
                    WHERE  charter_number = :cid
                    ORDER  BY period DESC
                    LIMIT  1
                """),
                {"cid": str(own_id)},
            ).fetchone()
        return dict(row._mapping) if row else {}
    return await asyncio.to_thread(_q, engine)


async def _fetch_recent_alerts(tenant_id: str, period: str, engine: sa.engine.Engine) -> list[dict]:
    def _q(eng):
        with eng.connect() as conn:
            rows = conn.execute(
                sa.text("""
                    SELECT alert_type, geography_label, subject_institution,
                           current_share, prior_share, change_pp, metric, narrative
                    FROM   alerts
                    WHERE  tenant_id = :tid AND period = :period
                    ORDER  BY change_pp DESC NULLS LAST
                    LIMIT  25
                """),
                {"tid": tenant_id, "period": period},
            ).fetchall()
        return [dict(r._mapping) for r in rows]
    return await asyncio.to_thread(_q, engine)


async def _fetch_peers(
    own_id: str, state: str, own_assets: float,
    period: str, engine: sa.engine.Engine,
) -> list[dict]:
    """Auto-select peers: same state, asset size within ±50%."""
    if not state or not own_assets:
        return []
    lo, hi = own_assets * 0.50, own_assets * 1.50
    def _q(eng):
        with eng.connect() as conn:
            rows = conn.execute(
                sa.text("""
                    SELECT DISTINCT ON (charter_number)
                           charter_number, institution_name, state,
                           total_assets, number_of_branches
                    FROM   institutions_quarterly
                    WHERE  state            = :state
                      AND  total_assets     BETWEEN :lo AND :hi
                      AND  charter_number  != :own
                    ORDER  BY charter_number, period DESC
                    LIMIT  :limit
                """),
                {"state": state, "lo": lo, "hi": hi, "own": str(own_id), "limit": _MAX_PEERS},
            ).fetchall()
        return [dict(r._mapping) for r in rows]
    return await asyncio.to_thread(_q, engine)

# ── Market-share computation ───────────────────────────────────────────────────

async def _geo_snapshot(
    geo: dict, period: str, own_id: str, engine: sa.engine.Engine,
) -> dict:
    """
    Pull deposit market share for one geography.
    Returns {geo, own_row, df, prior_period}.
    """
    df = await asyncio.to_thread(
        calculate_market_share,
        geo["geography_type"], geo["geography_id"], period,
        "deposits", ["credit_union", "bank"], engine,
    )
    own_row: Optional[dict] = None
    if df is not None and not df.empty:
        mask = df["charter_or_cert"].astype(str) == str(own_id)
        if mask.any():
            own_row = df[mask].iloc[0].to_dict()

    year, q = _parse_period(period)
    return {
        "geo":          geo,
        "own_id":       own_id,
        "own_row":      own_row,
        "df":           df if df is not None else pd.DataFrame(),
        "prior_period": _prior_quarter(year, q),
    }


async def _peer_metric_snapshot(
    primary_geo: dict,
    period: str,
    own_id: str,
    peer_ids: list[str],
    peer_branches: dict[str, int],
    engine: sa.engine.Engine,
) -> dict:
    """
    Returns {metrics: {key: {inst_id: value}}, institutions: [...], own_id}.
    Fetches the four share metrics from calculate_market_share in parallel.
    """
    all_ids = {own_id} | set(peer_ids)

    async def _pull(src_metric: str, out_key: str) -> dict[str, dict]:
        df = await asyncio.to_thread(
            calculate_market_share,
            primary_geo["geography_type"], primary_geo["geography_id"], period,
            src_metric, ["credit_union", "bank"], engine,
        )
        result: dict[str, dict] = {}
        if df is None or df.empty:
            return result
        for _, row in df.iterrows():
            cid = str(row["charter_or_cert"])
            if cid not in all_ids:
                continue
            result[cid] = {
                "share":      float(row["market_share"]),
                "change_pp":  float(row.get("share_change_prior_period") or 0),
                "confidence": str(row.get("confidence") or "estimated"),
                "name":       str(row.get("institution_name") or cid),
                "type":       str(row.get("institution_type") or ""),
            }
        return result

    dep, loan, mtg, mbr = await asyncio.gather(
        _pull("deposits",              "deposit_share"),
        _pull("loans",                 "loan_share"),
        _pull("mortgage_originations", "mortgage_share"),
        _pull("members",               "member_share"),
    )

    # Build flat institution list, own first
    seen: dict[str, dict] = {}
    for mapping in (dep, loan, mtg, mbr):
        for cid, d in mapping.items():
            seen.setdefault(cid, {"id": cid, "name": d["name"], "type": d["type"]})
    institutions = sorted(seen.values(), key=lambda x: (x["id"] != own_id, x["name"]))

    return {
        "own_id":       own_id,
        "institutions": institutions,
        "metrics": {
            "deposit_share":   {c: v["share"]     for c, v in dep.items()},
            "loan_share":      {c: v["share"]     for c, v in loan.items()},
            "mortgage_share":  {c: v["share"]     for c, v in mtg.items()},
            "member_share":    {c: v["share"]     for c, v in mbr.items()},
            "deposit_growth":  {c: v["change_pp"] for c, v in dep.items()},
            "total_branches":  {cid: float(peer_branches.get(cid, 0)) for cid in all_ids},
        },
        "confidence": {c: v["confidence"] for c, v in dep.items()},
    }

# ── Claude helpers ─────────────────────────────────────────────────────────────

_SYSTEM = (
    "You are the author of a quarterly competitive intelligence report for a credit union's "
    "strategic planning team. Write in a direct, data-driven tone suitable for a board of "
    "directors — authoritative, specific, and free of filler language. "
    "Prefer flowing paragraphs over bullet points. Never state 'it is important to note' "
    "or 'in conclusion'. Use specific numbers from the data provided."
)


async def _claude(prompt: str, max_tokens: int, semaphore: asyncio.Semaphore) -> str:
    async with semaphore:
        resp = await _client.messages.create(
            model=_MODEL,
            max_tokens=max_tokens,
            thinking={"type": "adaptive"},
            system=_SYSTEM,
            messages=[{"role": "user", "content": prompt}],
        )
    return next((b.text.strip() for b in resp.content if b.type == "text"), "")


def _fmt_pp(v: Optional[float]) -> str:
    if v is None:
        return "—"
    sign = "+" if v > 0 else ""
    return f"{sign}{v * 100:.1f}pp"

def _fmt_pct(v: Optional[float]) -> str:
    return f"{v * 100:.1f}%" if v is not None else "—"

def _median(vals: list[float]) -> Optional[float]:
    if not vals:
        return None
    s = sorted(vals)
    m = len(s) // 2
    return s[m] if len(s) % 2 else (s[m - 1] + s[m]) / 2


async def _write_exec_summary(
    cu_name: str, period_lbl: str,
    snapshots: list[dict], alerts: list[dict],
    peer_data: dict, semaphore: asyncio.Semaphore,
) -> str:
    own_id = peer_data.get("own_id", "")
    dep_shares = peer_data["metrics"].get("deposit_share", {})
    peer_ids = [i["id"] for i in peer_data["institutions"] if i["id"] != own_id]
    peer_vals = [dep_shares[pid] for pid in peer_ids if pid in dep_shares]
    own_dep   = dep_shares.get(own_id)
    median    = _median(peer_vals)

    # Aggregate own deposit share across monitored markets
    own_shares_by_geo = [
        f"{s['geo']['geography_label'] or s['geo']['geography_id']}: "
        f"{s['own_row']['market_share']*100:.1f}%"
        for s in snapshots if s.get("own_row")
    ]

    top_competitor = next(
        (a for a in alerts if a["alert_type"] == "competitor_gain"), None
    )
    top_opp = next(
        (a for a in alerts if a["alert_type"] in ("new_entrant", "market_growth")), None
    )

    prompt = (
        f"Write a 200-word (3 paragraphs, ~65 words each) executive summary for "
        f"{cu_name}'s competitive intelligence report covering {period_lbl}.\n\n"
        f"Deposit market share by monitored geography: "
        f"{'; '.join(own_shares_by_geo) or 'data unavailable'}.\n"
        f"Own share in primary market: {_fmt_pct(own_dep)}, "
        f"peer median: {_fmt_pct(median)}, "
        f"gap: {_fmt_pp((own_dep - median) if own_dep is not None and median is not None else None)}.\n"
        f"Most significant competitor movement: "
        f"{top_competitor['narrative'] if top_competitor else 'no significant competitor movements this quarter'}.\n"
        f"Top market opportunity signal: "
        f"{top_opp['narrative'] if top_opp else 'no unusual market growth signals detected'}.\n\n"
        "Paragraph 1: overall portfolio share position and trend. "
        "Paragraph 2: most significant competitive movement and its strategic implication. "
        "Paragraph 3: the single most actionable market opportunity."
    )
    return await _claude(prompt, 450, semaphore)


async def _write_geo_movement(
    cu_name: str, snapshot: dict, semaphore: asyncio.Semaphore,
) -> str:
    geo_label = snapshot["geo"]["geography_label"] or snapshot["geo"]["geography_id"]
    own       = snapshot.get("own_row")
    df        = snapshot.get("df", pd.DataFrame())
    own_id    = snapshot.get("own_id", "")

    if own is None or df.empty:
        return f"Detailed competitive data for {geo_label} was not available for this period."

    own_share  = own["market_share"]
    prior_pp   = own.get("share_change_prior_period") or 0
    yoy_pp     = own.get("share_change_yoy") or 0

    # Top 3 non-own institutions by share
    others = df[df["charter_or_cert"].astype(str) != str(own_id)].nlargest(3, "market_share")
    comp_lines = []
    for _, r in others.iterrows():
        chg = r.get("share_change_prior_period") or 0
        comp_lines.append(
            f"{r['institution_name']} ({r['market_share']*100:.1f}%, "
            f"{'+' if chg >= 0 else ''}{chg*100:.1f}pp QoQ)"
        )

    prompt = (
        f"Write 2–3 sentences analyzing competitive dynamics in {geo_label} "
        f"for {cu_name} in {snapshot.get('prior_period', 'the current period')}.\n\n"
        f"{cu_name} holds {own_share*100:.1f}% deposit share "
        f"({'+' if prior_pp >= 0 else ''}{prior_pp*100:.1f}pp vs prior quarter, "
        f"{'+' if yoy_pp >= 0 else ''}{yoy_pp*100:.1f}pp YoY).\n"
        f"Top competitors: {'; '.join(comp_lines) if comp_lines else 'no competitor data available'}.\n\n"
        "Identify the single most significant competitive movement, what it signals strategically, "
        "and whether it warrants a response. Do not restate every number — interpret the pattern."
    )
    return await _claude(prompt, 220, semaphore)


async def _write_opportunities(
    cu_name: str, snapshots: list[dict],
    peer_data: dict, semaphore: asyncio.Semaphore,
) -> str:
    own_id   = peer_data.get("own_id", "")
    dep_vals = peer_data["metrics"].get("deposit_share", {})
    peer_ids = [i["id"] for i in peer_data["institutions"] if i["id"] != own_id]
    peer_vals = [dep_vals[pid] for pid in peer_ids if pid in dep_vals]
    median = _median(peer_vals)
    own_primary = dep_vals.get(own_id)

    below_median_geos = []
    fast_growing_geos = []
    low_penetration_geos = []

    for s in snapshots:
        if not s.get("own_row"):
            continue
        lbl   = s["geo"]["geography_label"] or s["geo"]["geography_id"]
        share = s["own_row"]["market_share"]
        yoy   = s["own_row"].get("share_change_yoy") or 0

        if median and share < median * 0.92:
            below_median_geos.append(
                f"{lbl} (share {share*100:.1f}%, peer median ~{median*100:.1f}%)"
            )
        if yoy > 0.025:
            fast_growing_geos.append(f"{lbl} (+{yoy*100:.1f}pp YoY)")
        if share < 0.04:
            low_penetration_geos.append(f"{lbl} ({share*100:.1f}% share)")

    prompt = (
        f"Write 3 paragraphs identifying the top 3 market growth opportunities for {cu_name}. "
        "Each paragraph covers one opportunity geography.\n\n"
        f"Markets where {cu_name} significantly underperforms the peer median: "
        f"{', '.join(below_median_geos[:3]) or 'none identified'}.\n"
        f"Markets with rapid share gains (momentum to build on): "
        f"{', '.join(fast_growing_geos[:3]) or 'none'}.\n"
        f"Low-penetration markets (<4% deposit share) with room for growth: "
        f"{', '.join(low_penetration_geos[:3]) or 'none'}.\n\n"
        "For each opportunity: name the geography, explain the specific opportunity signal "
        "(underperformance vs peers, low CU penetration, competitor weakness, or growth momentum), "
        "and recommend the type of competitive response (e.g., branch expansion, CD rate promotion, "
        "mortgage campaign, small-business focus). Be concrete and specific."
    )
    return await _claude(prompt, 550, semaphore)


async def _write_peer_narrative(
    cu_name: str, peer_data: dict, semaphore: asyncio.Semaphore,
) -> str:
    own_id   = peer_data.get("own_id", "")
    metrics  = peer_data["metrics"]
    peer_ids = [i["id"] for i in peer_data["institutions"] if i["id"] != own_id]
    n_peers  = len(peer_ids)

    strengths, gaps = [], []
    for key, label, unit in _PEER_METRIC_ROWS[:5]:
        vals     = metrics.get(key, {})
        own_val  = vals.get(own_id)
        pvals    = [vals[pid] for pid in peer_ids if pid in vals]
        median   = _median(pvals)
        if own_val is None or median is None:
            continue
        diff = own_val - median
        if unit == "pct":
            txt = f"{label}: {own_val*100:.1f}% vs peer median {median*100:.1f}% ({_fmt_pp(diff)})"
        else:
            txt = f"{label}: {own_val*100:.1f}pp vs peer median {median*100:.1f}pp ({_fmt_pp(diff)})"
        (strengths if diff > 0 else gaps).append(txt)

    prompt = (
        f"Write 2 paragraphs benchmarking {cu_name} against its {n_peers} peer credit unions.\n\n"
        f"Areas of competitive strength (above peer median): "
        f"{'; '.join(strengths) or 'none identified'}.\n"
        f"Areas lagging the peer median: "
        f"{'; '.join(gaps) or 'none identified'}.\n\n"
        "Paragraph 1: competitive strengths — what drives them and how to protect the advantage. "
        "Paragraph 2: gaps — what they imply strategically and what targeted response is warranted. "
        "Be quantitative and interpretive, not just descriptive."
    )
    return await _claude(prompt, 380, semaphore)

# ── python-docx helpers ────────────────────────────────────────────────────────

def _shade_cell(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd   = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tc_pr.append(shd)


def _configure_doc(doc: Document) -> None:
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Normal"].font.color.rgb = _RGB_DARK
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)


def _add_cover(doc: Document, cu_name: str, period_lbl: str) -> None:
    for _ in range(7):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(cu_name)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = _RGB_BRAND

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Competitive Intelligence Report — {period_lbl}")
    r2.bold = True
    r2.font.size = Pt(18)
    r2.font.color.rgb = _RGB_DARK

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("CONFIDENTIAL — For Internal Use Only")
    r3.italic = True
    r3.font.size = Pt(11)
    r3.font.color.rgb = _RGB_RED

    doc.add_paragraph()

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(f"Prepared: {datetime.now().strftime('%B %Y')}")
    r4.font.size = Pt(10)
    r4.font.color.rgb = _RGB_MUTED


def _add_section_header(doc: Document, num: int, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(f"  {num}.  {title.upper()}")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = _RGB_WHITE
    # Shade the paragraph background blue via XML
    p_pr = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  _COL_HEADER)
    p_pr.append(shd)


def _add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(10)


def _add_note(doc: Document, text: str) -> None:
    p  = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = _RGB_MUTED


def _add_table_header(table, labels: list[str]) -> None:
    row = table.rows[0]
    for i, label in enumerate(labels):
        if i >= len(row.cells):
            break
        cell = row.cells[i]
        _shade_cell(cell, _COL_HEADER)
        cell.text = ""
        r = cell.paragraphs[0].add_run(label)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = _RGB_WHITE


def _cell_text(cell, text: str, size: int = 9) -> None:
    cell.text = str(text)
    for r in cell.paragraphs[0].runs:
        r.font.size = Pt(size)


def _add_market_share_table(
    doc: Document,
    snapshots: list[dict],
    own_id: str,
    peer_median: Optional[float],
) -> None:
    cols   = ["Geography", "Your Share", "vs Prior Qtr", "YoY Δ", "CU Rank", "vs Peer Median"]
    rows_n = min(len(snapshots), _MAX_GEO)
    table  = doc.add_table(rows=1 + rows_n, cols=len(cols))
    table.style = "Table Grid"
    _add_table_header(table, cols)

    for ri, snap in enumerate(snapshots[:rows_n]):
        row  = table.rows[ri + 1]
        geo  = snap["geo"]
        own  = snap.get("own_row")
        df   = snap.get("df", pd.DataFrame())
        lbl  = geo.get("geography_label") or geo["geography_id"]

        # Rank among CUs only
        cu_rank = "—"
        if not df.empty and own is not None:
            cu_df = (
                df[df["institution_type"] == "credit_union"]
                .sort_values("market_share", ascending=False)
                .reset_index(drop=True)
            )
            match = cu_df[cu_df["charter_or_cert"].astype(str) == str(own_id)]
            if not match.empty:
                cu_rank = f"#{match.index[0] + 1} of {len(cu_df)}"

        share    = own["market_share"]       if own else None
        prior_pp = (own.get("share_change_prior_period") or 0) if own else None
        yoy_pp   = (own.get("share_change_yoy") or 0)          if own else None
        vs_med   = (
            f"{_fmt_pp(share - peer_median)}"
            if (share is not None and peer_median is not None)
            else "—"
        )

        values = [
            lbl,
            _fmt_pct(share),
            _fmt_pp(prior_pp),
            _fmt_pp(yoy_pp),
            cu_rank,
            vs_med,
        ]
        for ci, val in enumerate(values):
            _cell_text(row.cells[ci], val)

        # Colour row by relative position vs peer median
        fill = _COL_ALT if ri % 2 == 0 else _WHITE
        if share is not None and peer_median is not None:
            if share > peer_median * 1.02:
                fill = _COL_GREEN
            elif share < peer_median * 0.98:
                fill = _COL_AMBER
        for ci in range(len(cols)):
            _shade_cell(row.cells[ci], fill)


def _add_peer_table(doc: Document, peer_data: dict) -> None:
    own_id   = peer_data.get("own_id", "")
    metrics  = peer_data["metrics"]
    peer_ids = [i["id"] for i in peer_data["institutions"] if i["id"] != own_id]
    all_ids  = [own_id] + peer_ids

    cols  = ["Metric", "You", "Peer Median", "Rank", "vs Median"]
    table = doc.add_table(rows=1 + len(_PEER_METRIC_ROWS), cols=len(cols))
    table.style = "Table Grid"
    _add_table_header(table, cols)

    for ri, (key, label, unit) in enumerate(_PEER_METRIC_ROWS):
        row      = table.rows[ri + 1]
        vals     = metrics.get(key, {})
        own_val  = vals.get(own_id)
        pvals    = [vals[pid] for pid in peer_ids if pid in vals]
        median   = _median(pvals)

        # Rank: 1 = highest, descending
        all_known = sorted([v for v in [own_val] + pvals if v is not None], reverse=True)
        rank = all_known.index(own_val) + 1 if own_val in all_known else None

        def fmt(v: Optional[float]) -> str:
            if v is None:
                return "—"
            if unit == "pct":
                return f"{v * 100:.1f}%"
            if unit == "pp":
                return f"{'+' if v > 0 else ''}{v * 100:.1f}pp"
            return str(int(round(v)))

        diff = (own_val - median) if (own_val is not None and median is not None) else None

        _cell_text(row.cells[0], label)
        _cell_text(row.cells[1], fmt(own_val))
        _cell_text(row.cells[2], fmt(median))
        _cell_text(row.cells[3], f"#{rank} of {len(all_known)}" if rank else "—")
        _cell_text(row.cells[4], fmt(diff) if diff is not None else "—")

        fill = _WHITE
        if diff is not None:
            fill = _COL_GREEN if diff > 0.001 else (_COL_AMBER if diff < -0.001 else _WHITE)
        elif ri % 2 == 0:
            fill = _COL_ALT
        for ci in range(len(cols)):
            _shade_cell(row.cells[ci], fill)


def _add_data_notes(doc: Document, snapshots: list[dict], period_lbl: str) -> None:
    # Confidence tally across all snapshot DataFrames
    conf_counts: dict[str, int] = {}
    for snap in snapshots:
        df = snap.get("df", pd.DataFrame())
        if df.empty or "confidence" not in df.columns:
            continue
        for c in df["confidence"].dropna():
            key = str(c).lower()
            conf_counts[key] = conf_counts.get(key, 0) + 1

    conf_summary = ", ".join(
        f"{v} {k}" for k, v in conf_counts.items() if v
    ) or "confidence levels not available"

    # Source table
    sources = [
        ("FDIC Summary of Deposits",        "Annual (June call date)",   "Deposit share — banks"),
        ("NCUA 5300 Call Report",            "Quarterly",                 "Deposit & loan share — credit unions"),
        ("HMDA Loan Application Register",   "Annual",                    "Mortgage origination market share"),
        ("U.S. Census ACS 5-Year Estimate",  "Annual",                    "Demographics, household income trends"),
    ]
    tbl = doc.add_table(rows=1 + len(sources), cols=3)
    tbl.style = "Table Grid"
    _add_table_header(tbl, ["Data Source", "Release Frequency", "Used For"])
    for ri, (src, freq, used) in enumerate(sources):
        r = tbl.rows[ri + 1]
        _cell_text(r.cells[0], src)
        _cell_text(r.cells[1], freq)
        _cell_text(r.cells[2], used)
        fill = _COL_ALT if ri % 2 == 0 else _WHITE
        for ci in range(3):
            _shade_cell(r.cells[ci], fill)

    doc.add_paragraph()
    _add_note(doc, f"Data quality for {period_lbl}: {conf_summary}.")

    # Confidence key
    doc.add_paragraph()
    defs = [
        ("Measured",  "Directly reported to the regulator (FDIC branch-level deposit call, HMDA origination record)."),
        ("Modeled",   "Branch-level deposit allocation model, validated against FDIC actuals; typical error < 5%."),
        ("Estimated", "Interpolated from annual data or derived from branch-count/demographic scaling."),
    ]
    for term, defn in defs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        bold_r = p.add_run(f"{term}: ")
        bold_r.bold = True
        bold_r.font.size = Pt(9)
        rest = p.add_run(defn)
        rest.font.size = Pt(9)

    # Next release dates
    doc.add_paragraph()
    hdr = doc.add_paragraph()
    hdr.add_run("Next scheduled data releases:").bold = True

    releases = [
        "FDIC Summary of Deposits — typically published October (for prior June call date)",
        "NCUA 5300 Call Report — released approximately 45 days after quarter-end",
        "HMDA Loan Application Register — typically published March (for prior calendar year)",
    ]
    for line in releases:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(line)
        r.font.size = Pt(9)

# ── Public entry point ────────────────────────────────────────────────────────

async def generate_quarterly_report(
    tenant_id: str,
    period: str,
    engine: Optional[sa.engine.Engine] = None,
) -> bytes:
    """
    Generate a full competitive intelligence Word report for one tenant.

    Parameters
    ----------
    tenant_id : str
        Identifies the credit union tenant.  Matches monitored_geographies.tenant_id.
    period : str
        Reporting period.  Accepts "2024Q4" or "2024" (treated as Q4).
    engine : SQLAlchemy Engine, optional
        Reuse an existing engine.  If omitted, one is created from DATABASE_URL.

    Returns
    -------
    bytes
        Raw .docx content, ready to write to disk or stream in an HTTP response.
    """
    if engine is None:
        engine = get_engine()

    year, q     = _parse_period(period)
    period_norm = _fmt_period(year, q)
    period_lbl  = _human_period(period_norm)

    semaphore = asyncio.Semaphore(5)

    # ── 1. Foundational data — parallel ───────────────────────────────────────
    geos, alerts = await asyncio.gather(
        _fetch_monitored_geos(tenant_id, engine),
        _fetch_recent_alerts(tenant_id, period_norm, engine),
    )

    # Own institution ID comes from the monitored geography configuration
    own_id = next(
        (g["own_institution_id"] for g in geos if g.get("own_institution_id")),
        tenant_id,
    )

    own_inst, = await asyncio.gather(
        _fetch_own_institution(own_id, period_norm, engine),
    )

    cu_name   = own_inst.get("institution_name") or tenant_id
    own_state = own_inst.get("state", "")
    own_assets = float(own_inst.get("total_assets") or 0)
    own_branches = int(own_inst.get("number_of_branches") or 0)

    peers = await _fetch_peers(own_id, own_state, own_assets, period_norm, engine)
    peer_ids = [str(p["charter_number"]) for p in peers]
    peer_branches = {str(p["charter_number"]): int(p.get("number_of_branches") or 0) for p in peers}
    peer_branches[own_id] = own_branches

    # ── 2. Per-geography market share snapshots — parallel ────────────────────
    snapshots = await asyncio.gather(
        *[_geo_snapshot(geo, period_norm, own_id, engine) for geo in geos]
    )

    # ── 3. Peer benchmarking data ──────────────────────────────────────────────
    primary_geo = geos[0] if geos else None
    if primary_geo and peer_ids:
        peer_data = await _peer_metric_snapshot(
            primary_geo, period_norm, own_id, peer_ids, peer_branches, engine
        )
    else:
        peer_data = {"own_id": own_id, "institutions": [], "metrics": {}, "confidence": {}}

    # Peer median deposit share (primary-market basis, used for colour-coding)
    dep_vals   = peer_data["metrics"].get("deposit_share", {})
    peer_dep_vals = [dep_vals[pid] for pid in peer_ids if pid in dep_vals]
    peer_median   = _median(peer_dep_vals)

    # ── 4. Generate Claude sections — all parallel, bounded by semaphore ──────
    exec_task  = asyncio.create_task(
        _write_exec_summary(cu_name, period_lbl, list(snapshots), alerts, peer_data, semaphore)
    )
    geo_tasks  = [
        asyncio.create_task(_write_geo_movement(cu_name, snap, semaphore))
        for snap in snapshots
    ]
    opp_task   = asyncio.create_task(
        _write_opportunities(cu_name, list(snapshots), peer_data, semaphore)
    )
    peer_task  = asyncio.create_task(
        _write_peer_narrative(cu_name, peer_data, semaphore)
    )

    exec_summary, *geo_narratives, opportunities, peer_narrative = await asyncio.gather(
        exec_task, *geo_tasks, opp_task, peer_task,
    )

    # ── 5. Assemble Word document ──────────────────────────────────────────────
    doc = Document()
    _configure_doc(doc)

    # Cover
    _add_cover(doc, cu_name, period_lbl)

    # ── Section 1: Executive Summary ──────────────────────────────────────────
    doc.add_page_break()
    _add_section_header(doc, 1, "Executive Summary")
    doc.add_paragraph()
    _add_body(doc, exec_summary)

    # ── Section 2: Market Share Dashboard ─────────────────────────────────────
    doc.add_page_break()
    _add_section_header(doc, 2, "Market Share Dashboard")
    doc.add_paragraph()
    _add_note(
        doc,
        f"Deposit market share by monitored geography — {period_lbl}.  "
        "Green = above peer median · Amber = below peer median · "
        "Share figures are fractions of total market deposits in that geography.",
    )
    doc.add_paragraph()
    _add_market_share_table(doc, list(snapshots), own_id, peer_median)

    # ── Section 3: Competitive Movements ──────────────────────────────────────
    doc.add_page_break()
    _add_section_header(doc, 3, "Competitive Movements")
    for snap, narrative in zip(snapshots, geo_narratives):
        lbl = snap["geo"].get("geography_label") or snap["geo"]["geography_id"]
        doc.add_paragraph()
        geo_hdr = doc.add_paragraph()
        geo_hdr.paragraph_format.space_after = Pt(4)
        r = geo_hdr.add_run(lbl)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = _RGB_BRAND
        _add_body(doc, narrative)

    # ── Section 4: Peer Benchmarking ──────────────────────────────────────────
    doc.add_page_break()
    _add_section_header(doc, 4, "Peer Benchmarking")
    doc.add_paragraph()
    n_peers = len(peer_ids)
    _add_note(
        doc,
        f"Comparison against {n_peers} auto-selected peer credit unions "
        f"(same state, ±50% asset size) in "
        f"{(geos[0].get('geography_label') or geos[0]['geography_id']) if geos else 'primary market'}, "
        f"{period_lbl}.  Green = above peer median · Amber = below peer median.",
    )
    doc.add_paragraph()
    _add_peer_table(doc, peer_data)
    doc.add_paragraph()
    _add_body(doc, peer_narrative)

    # ── Section 5: Market Opportunities ───────────────────────────────────────
    doc.add_page_break()
    _add_section_header(doc, 5, "Market Opportunities")
    doc.add_paragraph()
    _add_body(doc, opportunities)

    # ── Section 6: Data Notes ──────────────────────────────────────────────────
    doc.add_page_break()
    _add_section_header(doc, 6, "Data Notes")
    doc.add_paragraph()
    _add_data_notes(doc, list(snapshots), period_lbl)

    # ── Serialise to bytes ─────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
