"""
Credit quality report generator — python-docx output.

Public entry point:
    await generate_credit_quality_report(tenant_id, period, report_type) -> bytes

report_type values:
    "monthly_risk_committee"  — 5 sections, 4-6 pages
    "quarterly_board"         — adds trend charts, forward outlook, peer appendix

Data source: NCUA 5300 quarterly data (institutions_quarterly table).
All rates are stored as decimal fractions; this module converts to percentages at display.
Peer group = regional credit unions in the tenant's primary monitored geography.
"""
from __future__ import annotations

import asyncio
import io
import logging
from datetime import datetime
from typing import Any, Optional

import anthropic
import numpy as np
import sqlalchemy as sa
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from database import get_engine
from processing.delinquency_engine import get_regional_peers

logger = logging.getLogger(__name__)
_client = anthropic.AsyncAnthropic()

# ── Display config ────────────────────────────────────────────────────────────

_LOAN_TYPES: list[tuple[str, str]] = [
    ("delinq_rate_total",       "Total"),
    ("delinq_rate_real_estate", "Real Estate"),
    ("delinq_rate_auto",        "Auto"),
    ("delinq_rate_credit_card", "Credit Card"),
    ("delinq_rate_commercial",  "Commercial"),
    ("delinq_90plus_rate",      "90+ Days"),
]

_CO_METRICS: list[tuple[str, str]] = [
    ("chargeoff_rate_total",       "Total"),
    ("chargeoff_rate_auto",        "Auto"),
    ("chargeoff_rate_credit_card", "Credit Card"),
    ("chargeoff_rate_real_estate", "Real Estate"),
    ("chargeoff_rate_commercial",  "Commercial"),
]

_ALL_RATE_METRICS: list[str] = (
    [c for c, _ in _LOAN_TYPES]
    + [c for c, _ in _CO_METRICS]
    + ["alll_coverage_ratio"]
)

_TREND_FETCH_COLS: list[str] = _ALL_RATE_METRICS + ["institution_name", "oreo_balance"]

# ── Colors ────────────────────────────────────────────────────────────────────

_GREEN   = (0xD9, 0xEA, 0xD3)
_AMBER   = (0xFF, 0xF2, 0xCC)
_RED_C   = (0xF4, 0xCC, 0xCC)
_WHITE   = (0xFF, 0xFF, 0xFF)
_HDR_BG  = (0x1E, 0x3A, 0x8A)
_HDR_FG  = (0xFF, 0xFF, 0xFF)
_WARN_BG = (0xFE, 0xE2, 0xE2)
_WARN_FG = (0x99, 0x1B, 0x1B)
_OWN_HL  = (0xDB, 0xEA, 0xFE)   # light blue — own institution row highlight

# ── Format helpers ────────────────────────────────────────────────────────────

def _period_display(period: str) -> str:
    return f"Q{period[5]} {period[:4]}" if len(period) == 6 and "Q" in period else period

def _month_range(period: str) -> str:
    months = {"1": "January–March", "2": "April–June",
               "3": "July–September", "4": "October–December"}
    if "Q" in period:
        return f"{months.get(period[5], period[5:])} {period[:4]}"
    return period

def _prior_periods(period: str, n: int) -> list[str]:
    """Return list of n periods ending at period, newest first."""
    yr, q = int(period[:4]), int(period[5])
    result = [period]
    for _ in range(n - 1):
        q -= 1
        if q < 1:
            q, yr = 4, yr - 1
        result.append(f"{yr}Q{q}")
    return result

def _pct(v: Any, d: int = 2) -> str:
    return f"{float(v) * 100:.{d}f}%" if v is not None else "—"

def _ratio(v: Any, d: int = 2) -> str:
    return f"{float(v):.{d}f}×" if v is not None else "—"

def _chg_str(cur: Any, prv: Any) -> str:
    if cur is None or prv is None:
        return ""
    delta = float(cur) - float(prv)
    sign  = "+" if delta >= 0 else ""
    return f"{sign}{delta * 100:.2f}pp"

# ── DB helpers (sync — called via asyncio.to_thread) ─────────────────────────

def _fetch_tenant_info(tenant_id: str, engine: sa.engine.Engine) -> dict[str, Any]:
    with engine.connect() as conn:
        row = conn.execute(
            sa.text("""
                SELECT own_institution_id, geography_type, geography_id, geography_label
                FROM   monitored_geographies
                WHERE  tenant_id = :tid AND own_institution_id IS NOT NULL
                ORDER  BY id LIMIT 1
            """),
            {"tid": tenant_id},
        ).mappings().fetchone()
    if not row:
        raise ValueError(f"No monitored geography with institution found for tenant {tenant_id!r}")
    return dict(row)


def _fetch_institution_trend(
    charter: str, periods: list[str], cols: list[str], engine: sa.engine.Engine,
) -> dict[str, dict[str, Any]]:
    """Returns {period: {col: value, ...}} for all requested periods."""
    col_sql = ", ".join(cols)
    with engine.connect() as conn:
        rows = conn.execute(
            sa.text(f"""
                SELECT DISTINCT ON (data_period) data_period, {col_sql}
                FROM   institutions_quarterly
                WHERE  charter_number = :cn
                  AND  data_period    = ANY(:periods)
                ORDER  BY data_period, ingested_at DESC
            """),
            {"cn": charter, "periods": periods},
        ).mappings().fetchall()
    return {str(r["data_period"]): {c: r.get(c) for c in cols} for r in rows}


def _fetch_bulk_peer_aggregates(
    peers: list[str], periods: list[str], metrics: list[str],
    engine: sa.engine.Engine,
) -> dict[str, dict[str, Any]]:
    """
    Single query returning p10/p25/median/p75/p90/n for every metric × period.
    Returns {period: {metric_stat: value}}.
    """
    if not peers or not periods or not metrics:
        return {}

    exprs = []
    for col in metrics:
        for stat, frac in (("p10", 0.1), ("p25", 0.25), ("median", 0.5), ("p75", 0.75), ("p90", 0.9)):
            exprs.append(
                f"percentile_cont({frac}) WITHIN GROUP (ORDER BY {col}) AS {col}_{stat}"
            )
        exprs.append(f"COUNT({col}) AS {col}_n")

    sql = (
        f"SELECT data_period, {', '.join(exprs)} "
        f"FROM institutions_quarterly "
        f"WHERE charter_number = ANY(:peers) AND data_period = ANY(:periods) "
        f"GROUP BY data_period"
    )
    with engine.connect() as conn:
        rows = conn.execute(sa.text(sql), {"peers": peers, "periods": periods}).mappings().fetchall()
    return {str(r["data_period"]): {k: v for k, v in dict(r).items() if k != "data_period"}
            for r in rows}


def _fetch_regional_institutions(
    peers: list[str], period: str, engine: sa.engine.Engine,
) -> list[dict[str, Any]]:
    if not peers:
        return []
    with engine.connect() as conn:
        rows = conn.execute(
            sa.text("""
                SELECT DISTINCT ON (charter_number)
                       charter_number, institution_name,
                       delinq_rate_total
                FROM   institutions_quarterly
                WHERE  charter_number = ANY(:peers)
                  AND  data_period    = :period
                  AND  delinq_rate_total IS NOT NULL
                ORDER  BY charter_number, ingested_at DESC
            """),
            {"peers": peers, "period": period},
        ).mappings().fetchall()
    return sorted([dict(r) for r in rows],
                  key=lambda x: float(x["delinq_rate_total"]) if x["delinq_rate_total"] else 0)


def _fetch_p37_signals(
    geo_type: str, geo_id: str, engine: sa.engine.Engine,
) -> Optional[dict[str, Any]]:
    try:
        with engine.connect() as conn:
            row = conn.execute(
                sa.text("""
                    SELECT * FROM regional_economic_indicators
                    WHERE  geography_type = :gt AND geography_id = :gid
                    ORDER  BY period DESC LIMIT 1
                """),
                {"gt": geo_type, "gid": geo_id},
            ).mappings().fetchone()
        return dict(row) if row else None
    except Exception:
        return None

# ── Cell coloring ─────────────────────────────────────────────────────────────

def _rate_color(own: Any, median: Any, p75: Any) -> tuple[int, int, int]:
    """Green = below peer median; Amber = above median; Red = at or above P75."""
    if own is None or median is None:
        return _WHITE
    if p75 is not None and float(own) >= float(p75):
        return _RED_C
    if float(own) > float(median):
        return _AMBER
    return _GREEN


def _coverage_color(own: Any, median: Any) -> tuple[int, int, int]:
    """ALLL coverage: inverted — higher is better."""
    if own is None:
        return _WHITE
    if float(own) < 1.0:
        return _RED_C
    if median and float(own) < float(median):
        return _AMBER
    return _GREEN

# ── Document primitive helpers ────────────────────────────────────────────────

def _set_cell_bg(cell: Any, rgb: tuple[int, int, int]) -> None:
    hex_color = "{:02X}{:02X}{:02X}".format(*rgb)
    tc  = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _write_cell(
    cell: Any,
    lines: list[str],
    bg: Optional[tuple[int, int, int]] = None,
    center: bool = True,
    sizes: Optional[list[int]] = None,
    bold_first: bool = False,
    white_text: bool = False,
) -> None:
    """Clear a cell and write one or more lines with independent font sizes."""
    if bg:
        _set_cell_bg(cell, bg)

    # Remove all paragraphs beyond the first
    tc = cell._tc
    paras = tc.findall(qn("w:p"))
    for p in paras[1:]:
        tc.remove(p)

    def _make_run(para: Any, text: str, size: int, bold: bool) -> None:
        para.clear()
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        if white_text:
            run.font.color.rgb = RGBColor(*_HDR_FG)

    for i, text in enumerate(lines):
        size = sizes[i] if sizes and i < len(sizes) else (9 if i == 0 else 7)
        bold = bold_first and i == 0
        if i == 0:
            para = cell.paragraphs[0]
        else:
            para = cell.add_paragraph()
        _make_run(para, text, size, bold)
        if i > 0 and not white_text:
            para.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _style_doc(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin    = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)


def _add_section_heading(doc: Document, number: int, title: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"Section {number}: {title}")
    run.font.size  = Pt(14)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm  = OxmlElement("w:bottom")
    btm.set(qn("w:val"),   "single")
    btm.set(qn("w:sz"),    "6")
    btm.set(qn("w:space"), "1")
    btm.set(qn("w:color"), "1E3A8A")
    pBdr.append(btm)
    pPr.append(pBdr)
    doc.add_paragraph()


def _add_body_text(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)


def _add_red_callout(doc: Document, heading: str, body: str) -> None:
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, _WARN_BG)
    p0 = cell.paragraphs[0]
    p0.clear()
    r0 = p0.add_run(f"WARNING — {heading}")
    r0.font.bold  = True
    r0.font.size  = Pt(11)
    r0.font.color.rgb = RGBColor(*_WARN_FG)
    p1 = cell.add_paragraph(body)
    p1.runs[0].font.size = Pt(10)
    p1.runs[0].font.color.rgb = RGBColor(*_WARN_FG)
    doc.add_paragraph()


def _add_confidential_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CONFIDENTIAL — For Internal Use Only — Do Not Distribute")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

# ── Cover page ────────────────────────────────────────────────────────────────

def _add_cover_page(doc: Document, inst_name: str, period: str, report_type: str) -> None:
    for _ in range(7):
        doc.add_paragraph()

    def _center(text: str, size: int, bold: bool = False,
                 rgb: tuple = (0x1E, 0x3A, 0x8A)) -> None:
        p   = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size  = Pt(size)
        run.font.bold  = bold
        run.font.color.rgb = RGBColor(*rgb)

    _center(inst_name, 24, bold=True)
    _center("Credit Quality Report", 18)
    label = ("Monthly Risk Committee" if report_type == "monthly_risk_committee"
             else "Quarterly Board")
    _center(f"{label} | {_month_range(period)}", 13, rgb=(0x44, 0x44, 0x44))
    doc.add_paragraph()
    _center(f"Generated {datetime.now().strftime('%B %d, %Y')}", 10, rgb=(0x88, 0x88, 0x88))
    doc.add_paragraph()
    _center("CONFIDENTIAL", 18, bold=True, rgb=(0xDC, 0x26, 0x26))
    _center("For Internal Use Only — Do Not Distribute", 10, rgb=(0xDC, 0x26, 0x26))
    doc.add_page_break()

# ── Section 2: Delinquency trend table ───────────────────────────────────────

def _add_trend_table(
    doc: Document,
    periods: list[str],          # newest first
    inst_trend: dict[str, dict[str, Any]],
    peer_agg_all: dict[str, dict[str, Any]],
) -> None:
    display = list(reversed(periods))   # oldest → newest, left to right

    table = doc.add_table(rows=1 + len(_LOAN_TYPES), cols=1 + len(display))
    table.style = "Table Grid"

    # Header row
    _write_cell(table.cell(0, 0), ["Loan Type"],
                bg=_HDR_BG, bold_first=True, white_text=True, sizes=[9])
    for ci, p in enumerate(display):
        _write_cell(table.cell(0, ci + 1), [_period_display(p)],
                    bg=_HDR_BG, bold_first=True, white_text=True, sizes=[8])

    # Data rows
    for ri, (col, label) in enumerate(_LOAN_TYPES):
        row = ri + 1
        _write_cell(table.cell(row, 0), [label], bold_first=True, sizes=[9])
        for ci, p in enumerate(display):
            agg = peer_agg_all.get(p, {})
            own = inst_trend.get(p, {}).get(col)
            med = agg.get(f"{col}_median")
            p75 = agg.get(f"{col}_p75")
            color = _rate_color(own, med, p75)
            lines = [_pct(own)]
            if med is not None:
                lines.append(_pct(med))
            _write_cell(table.cell(row, ci + 1), lines, bg=color, sizes=[8, 7])

    p = doc.add_paragraph()
    run = p.add_run(
        "Cell: own rate (top) / peer median (bottom).  "
        "  Green = below peer median   Amber = above median   Red = at or above 75th percentile"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

# ── Section 3: Regional comparison table ─────────────────────────────────────

def _add_regional_table(
    doc: Document,
    own_name: str,
    own_charter: str,
    regional_insts: list[dict[str, Any]],
    period: str,
) -> None:
    if not regional_insts:
        doc.add_paragraph("No regional peer data available for this period.")
        return

    # Find own institution in the sorted list (sorted ascending by rate)
    own_rate = next(
        (float(r["delinq_rate_total"]) for r in regional_insts
         if str(r["charter_number"]) == str(own_charter)),
        None,
    )
    n = len(regional_insts)
    own_rank = next(
        (i + 1 for i, r in enumerate(regional_insts)
         if str(r["charter_number"]) == str(own_charter)),
        None,
    )

    median_rate = float(np.median([float(r["delinq_rate_total"]) for r in regional_insts]))

    p = doc.add_paragraph(
        f"{n} credit unions with delinquency data for {_period_display(period)} | "
        f"Regional median: {median_rate * 100:.2f}%"
        + (f" | {own_name}: {own_rate * 100:.2f}% (rank #{own_rank} of {n})"
           if own_rate and own_rank else "")
    )
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_paragraph()

    # Show top 15 and bottom 5 to keep the table short; always include own institution
    def _include(i: int, r: dict) -> bool:
        is_own = str(r["charter_number"]) == str(own_charter)
        return is_own or i < 10 or i >= n - 3

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for ci, hdr in enumerate(["Institution", "Delinq Rate", "vs. Regional Median"]):
        _write_cell(table.cell(0, ci), [hdr], bg=_HDR_BG, bold_first=True, white_text=True, sizes=[9])

    prev_included = True
    for i, inst in enumerate(regional_insts):
        if not _include(i, inst):
            if prev_included:
                # Ellipsis row
                row = table.add_row()
                _write_cell(row.cells[0], ["…"], center=False, sizes=[9])
                _write_cell(row.cells[1], [""], sizes=[9])
                _write_cell(row.cells[2], [""], sizes=[9])
            prev_included = False
            continue
        prev_included = True

        is_own = str(inst["charter_number"]) == str(own_charter)
        rate   = float(inst["delinq_rate_total"])
        diff   = rate - median_rate
        diff_str = f"{'+'  if diff >= 0 else ''}{diff * 100:.2f}pp"
        name_disp = own_name if is_own else f"Peer {i + 1:02d}"
        bg    = _OWN_HL if is_own else _WHITE
        row = table.add_row()
        _write_cell(row.cells[0], [name_disp + (" ◀" if is_own else "")],
                    bg=bg, center=False, bold_first=is_own, sizes=[9])
        _write_cell(row.cells[1], [_pct(rate)], bg=bg, sizes=[9])
        color = _AMBER if diff > 0 else _GREEN
        _write_cell(row.cells[2], [diff_str], bg=color if not is_own else bg, sizes=[9])

    doc.add_paragraph()

# ── Section 4: Charge-off and coverage summary ────────────────────────────────

def _add_chargeoff_coverage_table(
    doc: Document,
    own_cur: dict[str, Any],
    own_prv: dict[str, Any],
    inst_trend: dict[str, dict[str, Any]],
    period: str,
    peer_agg_all: dict[str, dict[str, Any]],
    periods: list[str],
) -> None:
    peer_cur  = peer_agg_all.get(period, {})
    prior_per = periods[1] if len(periods) > 1 else None
    year_ago  = periods[4] if len(periods) > 4 else None

    # ── Charge-off sub-table ──
    p = doc.add_paragraph("Charge-Off Rates (annualized)")
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(11)

    tbl = doc.add_table(rows=1 + len(_CO_METRICS), cols=4)
    tbl.style = "Table Grid"
    for ci, hdr in enumerate(["Loan Type", "Current Quarter", "Prior Quarter", "Peer Median"]):
        _write_cell(tbl.cell(0, ci), [hdr], bg=_HDR_BG, bold_first=True, white_text=True, sizes=[9])

    for ri, (col, label) in enumerate(_CO_METRICS):
        row = ri + 1
        cur  = own_cur.get(col)
        prv  = own_prv.get(col) if own_prv else None
        pmed = peer_cur.get(f"{col}_median")
        _write_cell(tbl.cell(row, 0), [label], center=False, sizes=[9])
        _write_cell(tbl.cell(row, 1), [_pct(cur)], sizes=[9])
        chg_line = [_pct(prv), _chg_str(cur, prv)] if prv else [_pct(prv)]
        _write_cell(tbl.cell(row, 2), chg_line, sizes=[9, 7])
        color = _rate_color(cur, pmed, peer_cur.get(f"{col}_p75"))
        _write_cell(tbl.cell(row, 3), [_pct(pmed)], bg=color, sizes=[9])

    doc.add_paragraph()

    # ── ALLL coverage sub-table ──
    p2 = doc.add_paragraph("ALLL Coverage Ratio")
    p2.runs[0].font.bold = True
    p2.runs[0].font.size = Pt(11)

    cur_alll   = own_cur.get("alll_coverage_ratio")
    prv_alll   = own_prv.get("alll_coverage_ratio") if own_prv else None
    year_alll  = inst_trend.get(year_ago, {}).get("alll_coverage_ratio") if year_ago else None
    pmed_alll  = peer_cur.get("alll_coverage_ratio_median")

    tbl2 = doc.add_table(rows=2, cols=4)
    tbl2.style = "Table Grid"
    for ci, hdr in enumerate(["Metric", "Current", "Prior Year", "Peer Median"]):
        _write_cell(tbl2.cell(0, ci), [hdr], bg=_HDR_BG, bold_first=True, white_text=True, sizes=[9])

    alll_color = _coverage_color(cur_alll, pmed_alll)
    _write_cell(tbl2.cell(1, 0), ["ALLL Coverage"], center=False, sizes=[9])
    _write_cell(tbl2.cell(1, 1), [_ratio(cur_alll)], bg=alll_color, sizes=[9])
    yr_lines = [_ratio(year_alll)]
    if year_alll and cur_alll:
        yr_lines.append(_chg_str(cur_alll, year_alll))
    _write_cell(tbl2.cell(1, 2), yr_lines, sizes=[9, 7])
    _write_cell(tbl2.cell(1, 3), [_ratio(pmed_alll)], sizes=[9])

    doc.add_paragraph()

# ── Section 5: Watch items ────────────────────────────────────────────────────

def _add_watch_items(doc: Document, items: list[dict[str, str]]) -> None:
    if not items:
        doc.add_paragraph("No specific watch items identified for this period.")
        return
    for i, item in enumerate(items, 1):
        tbl  = doc.add_table(rows=4, cols=2)
        tbl.style = "Table Grid"
        labels = ["Metric", "Current Value", "Peer Context", "Recommended Action"]
        keys   = ["item", "value", "context", "action"]
        for ri, (lbl, key) in enumerate(zip(labels, keys)):
            _write_cell(tbl.cell(ri, 0), [lbl], bg=_HDR_BG, bold_first=True,
                        white_text=True, center=False, sizes=[9])
            val = item.get(key, "")
            bold = key == "action"
            color = _WARN_BG if key == "action" else _WHITE
            _write_cell(tbl.cell(ri, 1), [val], bg=color, center=False,
                        bold_first=bold, sizes=[9])
        doc.add_paragraph()

# ── Charts (quarterly board only) ────────────────────────────────────────────

def _render_trend_chart(
    periods: list[str],
    own_vals: list[Optional[float]],
    peer_meds: list[Optional[float]],
    p25s: list[Optional[float]],
    p75s: list[Optional[float]],
    title: str,
) -> io.BytesIO:
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(7, 3.2), dpi=150)
    x = list(range(len(periods)))
    xlabels = [_period_display(p) for p in periods]

    def _f(lst: list) -> list[Optional[float]]:
        return [float(v) * 100 if v is not None else None for v in lst]

    ov, pm, lo, hi = _f(own_vals), _f(peer_meds), _f(p25s), _f(p75s)

    # Band
    xi = [i for i, (a, b) in enumerate(zip(lo, hi)) if a is not None and b is not None]
    if xi:
        ax.fill_between(xi, [lo[i] for i in xi], [hi[i] for i in xi],
                        alpha=0.12, color="#2563eb", label="Peer P25–P75 band")

    # Peer median dashed
    xm = [i for i, v in enumerate(pm) if v is not None]
    if xm:
        ax.plot(xm, [pm[i] for i in xm], color="#94a3b8", linewidth=1.5,
                linestyle="--", label="Peer median")

    # Own institution solid
    xo = [i for i, v in enumerate(ov) if v is not None]
    if xo:
        ax.plot(xo, [ov[i] for i in xo], color="#2563eb", linewidth=2.2,
                marker="o", markersize=4, label="Your institution")

    ax.set_xticks(x)
    ax.set_xticklabels(xlabels, rotation=28, ha="right", fontsize=7)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v, _: f"{v:.1f}%"))
    ax.set_title(title, fontsize=10, pad=8)
    ax.grid(axis="y", alpha=0.25)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.legend(fontsize=7, loc="upper left")
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    buf.seek(0)
    plt.close(fig)
    return buf


def _render_loan_type_chart(
    labels: list[str],
    own_rates: list[float],
    peer_meds: list[float],
    title: str,
) -> io.BytesIO:
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(7, 3.2), dpi=150)
    x = np.arange(len(labels))
    w = 0.35
    ax.bar(x - w / 2, [v * 100 for v in own_rates],  w, label="Your institution",
           color="#2563eb", alpha=0.85)
    ax.bar(x + w / 2, [v * 100 for v in peer_meds], w, label="Peer median",
           color="#94a3b8", alpha=0.85)
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=9)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v, _: f"{v:.1f}%"))
    ax.set_title(title, fontsize=10, pad=8)
    ax.grid(axis="y", alpha=0.25)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.legend(fontsize=8)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    buf.seek(0)
    plt.close(fig)
    return buf


def _add_trend_charts(
    doc: Document,
    periods: list[str],             # newest first
    inst_trend: dict[str, dict[str, Any]],
    peer_agg_all: dict[str, dict[str, Any]],
) -> None:
    oldest_first = list(reversed(periods))
    period = periods[0]
    peer_cur = peer_agg_all.get(period, {})

    def _series(col: str) -> tuple:
        own_v = [inst_trend.get(p, {}).get(col) for p in oldest_first]
        meds  = [peer_agg_all.get(p, {}).get(f"{col}_median") for p in oldest_first]
        p25s  = [peer_agg_all.get(p, {}).get(f"{col}_p25") for p in oldest_first]
        p75s  = [peer_agg_all.get(p, {}).get(f"{col}_p75") for p in oldest_first]
        return own_v, meds, p25s, p75s

    # Chart 1 — total delinquency trend
    ov, meds, lo, hi = _series("delinq_rate_total")
    buf1 = _render_trend_chart(oldest_first, ov, meds, lo, hi,
                               "Total Delinquency Rate — 8 Quarter Trend")
    doc.add_picture(buf1, width=Inches(6.2))
    doc.add_paragraph()

    # Chart 2 — loan type comparison for current period
    lt_labels = [lbl for _, lbl in _LOAN_TYPES]
    lt_own    = [float(inst_trend.get(period, {}).get(col) or 0) for col, _ in _LOAN_TYPES]
    lt_peer   = [float(peer_cur.get(f"{col}_median") or 0) for col, _ in _LOAN_TYPES]
    buf2 = _render_loan_type_chart(lt_labels, lt_own, lt_peer,
                                   f"Delinquency by Loan Type — {_period_display(period)}")
    doc.add_picture(buf2, width=Inches(6.2))
    doc.add_paragraph()

    # Chart 3 — charge-off trend
    ov_co, meds_co, lo_co, hi_co = _series("chargeoff_rate_total")
    buf3 = _render_trend_chart(oldest_first, ov_co, meds_co, lo_co, hi_co,
                               "Annualized Net Charge-Off Rate — 8 Quarter Trend")
    doc.add_picture(buf3, width=Inches(6.2))
    doc.add_paragraph()

# ── Peer distribution appendix ────────────────────────────────────────────────

def _add_peer_distribution_appendix(
    doc: Document,
    peer_agg: dict[str, Any],
    own_metrics: dict[str, Any],
) -> None:
    all_metrics = [
        ("delinq_rate_total",       "Total Delinquency"),
        ("delinq_rate_auto",        "Auto Delinquency"),
        ("delinq_rate_real_estate", "Real Estate Delinquency"),
        ("delinq_rate_credit_card", "Credit Card Delinquency"),
        ("delinq_rate_commercial",  "Commercial Delinquency"),
        ("delinq_90plus_rate",      "90+ Day Delinquency"),
        ("chargeoff_rate_total",    "Net Charge-Off Rate (annualized)"),
        ("alll_coverage_ratio",     "ALLL Coverage Ratio"),
    ]

    tbl = doc.add_table(rows=1 + len(all_metrics), cols=7)
    tbl.style = "Table Grid"
    for ci, hdr in enumerate(["Metric", "P10", "P25", "Median", "P75", "P90", "Your Rate"]):
        _write_cell(tbl.cell(0, ci), [hdr], bg=_HDR_BG, bold_first=True,
                    white_text=True, sizes=[8])

    for ri, (col, label) in enumerate(all_metrics):
        row = ri + 1
        is_coverage = col == "alll_coverage_ratio"
        fmt = _ratio if is_coverage else _pct

        p10  = peer_agg.get(f"{col}_p10")
        p25  = peer_agg.get(f"{col}_p25")
        med  = peer_agg.get(f"{col}_median")
        p75  = peer_agg.get(f"{col}_p75")
        p90  = peer_agg.get(f"{col}_p90")
        own  = own_metrics.get(col)

        color = (_coverage_color(own, med) if is_coverage
                 else _rate_color(own, med, p75))

        _write_cell(tbl.cell(row, 0), [label], center=False, sizes=[8])
        for ci, v in enumerate([p10, p25, med, p75, p90], 1):
            _write_cell(tbl.cell(row, ci), [fmt(v)], sizes=[8])
        _write_cell(tbl.cell(row, 6), [fmt(own)], bg=color, bold_first=True, sizes=[8])

    doc.add_paragraph()

# ── Claude prompt functions ───────────────────────────────────────────────────

async def _claude_executive_summary(
    inst_name: str, period_label: str,
    own: dict[str, Any], peer: dict[str, Any],
) -> str:
    own_total  = own.get("delinq_rate_total")
    peer_total = peer.get("delinq_rate_total_median")
    comparison = ""
    if own_total and peer_total:
        rel = "above" if float(own_total) > float(peer_total) else (
              "below" if float(own_total) < float(peer_total) else "at")
        comparison = f"({rel} peer median of {float(peer_total) * 100:.2f}%)"

    lines = [
        f"Institution: {inst_name}",
        f"Period: {period_label}",
        f"Total delinquency: {_pct(own_total)} {comparison}",
    ]
    for col, lbl in _LOAN_TYPES[1:]:
        v = own.get(col)
        pm = peer.get(f"{col}_median")
        if v:
            lines.append(f"  {lbl}: {_pct(v)}" + (f" (peer: {_pct(pm)})" if pm else ""))
    alll = own.get("alll_coverage_ratio")
    if alll:
        lines.append(f"ALLL coverage: {_ratio(alll)} (minimum adequate: 1.0×)")
    co = own.get("chargeoff_rate_total")
    if co:
        lines.append(f"Net charge-off rate (annualized): {_pct(co)}")

    response = await _client.messages.create(
        model="claude-opus-4-8",
        max_tokens=350,
        thinking={"type": "adaptive"},
        system=(
            "You write 150-word (±10) credit quality executive summaries for risk committees. "
            "Structure (do not label sections): "
            "(1) State the total delinquency rate and whether it is above, below, or at the peer median. "
            "(2) Name the single most significant loan-type movement with its specific rate. "
            "(3) Name one specific concern with the metric that drives it. "
            "(4) Name one genuine positive — use the data, not a generic observation. "
            "Use exact numbers throughout. No hedging. Do not start with the institution name or 'As of'."
        ),
        messages=[{"role": "user", "content": "\n".join(lines)}],
    )
    blocks = [b for b in response.content if b.type == "text"]
    return blocks[-1].text.strip() if blocks else "(Summary unavailable)"


async def _claude_regional_narrative(
    inst_name: str, period_label: str,
    own_rate: Any,
    regional_insts: list[dict[str, Any]],
    p37: Optional[dict[str, Any]],
) -> str:
    rates = sorted(
        [float(r["delinq_rate_total"]) for r in regional_insts if r.get("delinq_rate_total")]
    )
    regional_median = float(np.median(rates)) if rates else None
    n_above = sum(1 for r in rates if r > float(own_rate or 0))
    n_total = len(rates)
    pct_rank = round((1 - n_above / n_total) * 100) if n_total else None

    lines = [
        f"Institution: {inst_name}",
        f"Period: {period_label}",
        f"Own total delinquency: {float(own_rate or 0) * 100:.2f}%",
        f"Regional median: {regional_median * 100:.2f}%" if regional_median else "",
        f"Institutions with higher delinquency: {n_above} of {n_total}",
        f"Own institution percentile rank: {pct_rank}th percentile" if pct_rank else "",
    ]
    if p37:
        for key in ("employer_hiring_signal", "business_closure_rate",
                    "permit_activity", "unemployment_rate"):
            v = p37.get(key)
            if v:
                lines.append(f"Regional economic signal — {key.replace('_', ' ')}: {v}")

    response = await _client.messages.create(
        model="claude-opus-4-8",
        max_tokens=280,
        thinking={"type": "adaptive"},
        system=(
            "You are a credit analyst writing a 2–3 sentence geographic context paragraph "
            "for a risk committee report. "
            "Answer: is this institution's delinquency driven by regional economic conditions "
            "(other institutions in the area are also elevated) or is it institution-specific? "
            "State your conclusion in the first sentence with the percentile rank as evidence. "
            "Use the P37 economic signals to confirm or challenge the conclusion. "
            "Be direct — give a definitive interpretation, not a hedge."
        ),
        messages=[{"role": "user", "content": "\n".join(l for l in lines if l)}],
    )
    blocks = [b for b in response.content if b.type == "text"]
    return blocks[-1].text.strip() if blocks else "(Analysis unavailable)"


async def _claude_watch_items(
    inst_name: str, period_label: str,
    own: dict[str, Any], peer: dict[str, Any], own_prv: dict[str, Any],
) -> list[dict[str, str]]:
    lines = [f"Institution: {inst_name}", f"Period: {period_label}", ""]
    for col, lbl in _LOAN_TYPES:
        v   = own.get(col)
        pm  = peer.get(f"{col}_median")
        prv = own_prv.get(col)
        if v is None:
            continue
        line = f"{lbl} delinquency: {float(v) * 100:.2f}%"
        if pm:
            rel = "above" if float(v) > float(pm) else "below"
            line += f" (peer median {float(pm) * 100:.2f}%, {rel})"
        if prv is not None:
            line += f" | QoQ: {(float(v) - float(prv)) * 100:+.2f}pp"
        lines.append(line)

    alll = own.get("alll_coverage_ratio")
    if alll:
        lines.append(f"ALLL coverage: {float(alll):.2f}× (minimum: 1.0×)")
    co = own.get("chargeoff_rate_total")
    co_prv = own_prv.get("chargeoff_rate_total")
    if co:
        line = f"Annualized net charge-off rate: {float(co) * 100:.2f}%"
        if co_prv and float(co_prv) > 0:
            accel = (float(co) - float(co_prv)) / float(co_prv) * 100
            line += f" | QoQ acceleration: {accel:+.0f}%"
        lines.append(line)

    response = await _client.messages.create(
        model="claude-opus-4-8",
        max_tokens=700,
        thinking={"type": "adaptive"},
        system=(
            "You are a senior credit analyst preparing a risk committee briefing. "
            "Identify the top 3 items requiring management attention. "
            "For each, respond in this exact format (no other text):\n\n"
            "ITEM: [specific metric or issue name]\n"
            "VALUE: [current value with units]\n"
            "CONTEXT: [one sentence peer comparison or trend context with specific numbers]\n"
            "ACTION: [one specific action management should take this quarter]\n"
            "---\n\n"
            "Rank by severity. Every item must cite at least one specific number."
        ),
        messages=[{"role": "user", "content": "\n".join(lines)}],
    )
    blocks = [b for b in response.content if b.type == "text"]
    raw    = blocks[-1].text.strip() if blocks else ""

    items: list[dict[str, str]] = []
    for block in raw.split("---"):
        block = block.strip()
        if not block:
            continue
        item: dict[str, str] = {}
        for line in block.splitlines():
            for key in ("ITEM", "VALUE", "CONTEXT", "ACTION"):
                if line.upper().startswith(f"{key}:"):
                    item[key.lower()] = line[len(key) + 1:].strip()
        if "item" in item:
            items.append(item)
    return items[:3]


async def _claude_forward_outlook(
    inst_name: str, period_label: str,
    inst_trend: dict[str, dict[str, Any]],
    periods: list[str],
    peer_agg_all: dict[str, dict[str, Any]],
    p37: Optional[dict[str, Any]],
) -> str:
    lines = [
        f"Institution: {inst_name}",
        f"Most recent period: {period_label}",
        "",
        "Total delinquency trend (newest → oldest):",
    ]
    for p in periods[:6]:
        v   = inst_trend.get(p, {}).get("delinq_rate_total")
        pm  = peer_agg_all.get(p, {}).get("delinq_rate_total_median")
        if v is not None:
            line = f"  {_period_display(p)}: {float(v) * 100:.2f}%"
            if pm:
                line += f" (peer median: {float(pm) * 100:.2f}%)"
            lines.append(line)

    if p37:
        lines.append("")
        lines.append("Regional economic indicators:")
        for key in ("employer_hiring_signal", "business_closure_rate",
                    "permit_activity", "unemployment_rate", "consumer_confidence"):
            v = p37.get(key)
            if v:
                lines.append(f"  {key.replace('_', ' ').title()}: {v}")

    response = await _client.messages.create(
        model="claude-opus-4-8",
        max_tokens=450,
        thinking={"type": "adaptive"},
        system=(
            "You are a senior credit economist writing a forward-looking outlook section "
            "for a credit union board report. "
            "Write 3–4 sentences describing the most probable credit quality scenario "
            "for the next two quarters. "
            "Base the projection on: (1) the trend trajectory, (2) the gap between the "
            "institution and peer median, and (3) regional economic signals. "
            "Include specific numeric ranges for expected delinquency. "
            "Note the primary downside risk and one mitigating factor. "
            "This is a professional board communication — be direct and precise."
        ),
        messages=[{"role": "user", "content": "\n".join(lines)}],
    )
    blocks = [b for b in response.content if b.type == "text"]
    return blocks[-1].text.strip() if blocks else "(Outlook unavailable)"

# ── Main entry point ──────────────────────────────────────────────────────────

async def generate_credit_quality_report(
    tenant_id: str,
    period: str,
    report_type: str,
    engine: Optional[sa.engine.Engine] = None,
) -> bytes:
    """
    Build a credit quality report and return the .docx file as bytes.

    Parameters
    ----------
    tenant_id   : str  — tenant JWT claim; used to resolve institution + geography
    period      : str  — e.g. "2024Q4"
    report_type : str  — "monthly_risk_committee" | "quarterly_board"

    Returns
    -------
    bytes  — raw .docx content; caller writes to disk or returns as HTTP response
    """
    if report_type not in ("monthly_risk_committee", "quarterly_board"):
        raise ValueError(
            f"report_type must be 'monthly_risk_committee' or 'quarterly_board', "
            f"got {report_type!r}"
        )

    if engine is None:
        engine = get_engine()

    # ── 1. Resolve tenant context ─────────────────────────────────────────────
    info = await asyncio.to_thread(_fetch_tenant_info, tenant_id, engine)
    charter   = info["own_institution_id"]
    geo_type  = info["geography_type"]
    geo_id    = info["geography_id"]
    geo_label = info.get("geography_label") or geo_id

    periods_8q  = _prior_periods(period, 8)   # newest first
    prior_period = periods_8q[1] if len(periods_8q) > 1 else None

    # ── 2. Fetch all institution data + peer group in parallel ─────────────────
    inst_trend_task = asyncio.to_thread(
        _fetch_institution_trend, charter, periods_8q, _TREND_FETCH_COLS, engine
    )
    peers_task = asyncio.to_thread(
        get_regional_peers, charter, geo_type, geo_id, period, engine
    )
    p37_task = asyncio.to_thread(_fetch_p37_signals, geo_type, geo_id, engine)

    inst_trend, peers, p37 = await asyncio.gather(inst_trend_task, peers_task, p37_task)

    inst_name    = inst_trend.get(period, {}).get("institution_name") or charter
    own_metrics  = inst_trend.get(period, {})
    own_metrics_prior = inst_trend.get(prior_period, {}) if prior_period else {}

    # ── 3. Bulk peer aggregates (single query) ─────────────────────────────────
    peer_agg_all = await asyncio.to_thread(
        _fetch_bulk_peer_aggregates, peers, periods_8q, _ALL_RATE_METRICS, engine
    )
    peer_agg_current = peer_agg_all.get(period, {})

    # ── 4. Regional institutions for Section 3 ────────────────────────────────
    regional_insts = await asyncio.to_thread(
        _fetch_regional_institutions, peers, period, engine
    )

    # ── 5. All Claude calls in parallel ───────────────────────────────────────
    claude_coros = [
        _claude_executive_summary(inst_name, _period_display(period),
                                  own_metrics, peer_agg_current),
        _claude_regional_narrative(inst_name, _period_display(period),
                                   own_metrics.get("delinq_rate_total"),
                                   regional_insts, p37),
        _claude_watch_items(inst_name, _period_display(period),
                            own_metrics, peer_agg_current, own_metrics_prior),
    ]
    if report_type == "quarterly_board":
        claude_coros.append(
            _claude_forward_outlook(inst_name, _period_display(period),
                                    inst_trend, periods_8q, peer_agg_all, p37)
        )

    claude_results = await asyncio.gather(*claude_coros)
    exec_summary      = claude_results[0]
    regional_narrative = claude_results[1]
    watch_items        = claude_results[2]
    forward_outlook    = claude_results[3] if report_type == "quarterly_board" else None

    # ── 6. Assemble document ──────────────────────────────────────────────────
    doc = Document()
    _style_doc(doc)
    _add_cover_page(doc, inst_name, period, report_type)

    # Section 1: Executive Summary
    _add_section_heading(doc, 1, "Executive Summary")
    _add_body_text(doc, exec_summary)
    doc.add_paragraph()

    # Section 2: Delinquency Trend
    _add_section_heading(doc, 2, "Delinquency Trend — 8 Quarter History")
    _add_trend_table(doc, periods_8q, inst_trend, peer_agg_all)

    # Section 3: Regional Comparison
    _add_section_heading(doc, 3, f"Regional Comparison — {geo_label}")
    _add_regional_table(doc, inst_name, charter, regional_insts, period)
    _add_body_text(doc, regional_narrative)
    doc.add_paragraph()

    # Section 4: Charge-Off and Coverage Summary
    _add_section_heading(doc, 4, "Charge-Off and Coverage Summary")
    alll_cov = own_metrics.get("alll_coverage_ratio")
    if alll_cov is not None and float(alll_cov) < 1.0:
        peer_alll_med = peer_agg_current.get("alll_coverage_ratio_median")
        gap = 1.0 - float(alll_cov)
        _add_red_callout(
            doc,
            "Inadequate ALLL Coverage",
            (
                f"Current coverage is {float(alll_cov):.2f}× — {gap:.2f}× below the 1.0× examiner "
                "minimum. Reserves are insufficient to cover the full delinquent balance. "
                + (f"Peer median coverage is {float(peer_alll_med):.2f}×. " if peer_alll_med else "")
                + "Recommend immediate ALLL adequacy review and additional provisioning "
                "before the next regulatory examination."
            ),
        )
    _add_chargeoff_coverage_table(doc, own_metrics, own_metrics_prior,
                                  inst_trend, period, peer_agg_all, periods_8q)

    # Section 5: Watch Items
    _add_section_heading(doc, 5, "Watch Items — Management Attention Required")
    _add_watch_items(doc, watch_items)

    # ── Quarterly board additions ─────────────────────────────────────────────
    if report_type == "quarterly_board":
        doc.add_page_break()

        _add_section_heading(doc, 6, "Year-over-Year Trend Charts")
        _add_trend_charts(doc, periods_8q, inst_trend, peer_agg_all)

        _add_section_heading(doc, 7, "Forward Outlook — Next Two Quarters")
        _add_body_text(doc, forward_outlook or "(Outlook unavailable)")
        doc.add_paragraph()
        p_note = doc.add_paragraph(
            "Note: This outlook is based on current delinquency trajectories and available regional "
            "economic indicators. It does not constitute a guarantee of future performance."
        )
        p_note.runs[0].font.size = Pt(9)
        p_note.runs[0].font.color.rgb = RGBColor(0x77, 0x77, 0x77)
        p_note.runs[0].font.italic = True

        doc.add_page_break()
        _add_section_heading(doc, 8, "Appendix — Anonymized Peer Distribution")
        p_note2 = doc.add_paragraph(
            f"Based on {len(peers)} regional credit unions with data for "
            f"{_period_display(period)}. Peer institutions are not identified by name."
        )
        p_note2.runs[0].font.size = Pt(9)
        p_note2.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        doc.add_paragraph()
        _add_peer_distribution_appendix(doc, peer_agg_current, own_metrics)

    _add_confidential_footer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
